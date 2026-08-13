"""
Generate FeedWise UNESCO Project Documentation (unescoproject.docx)
Run: python3 /home/upendra-singh-dhami/Documents/feedwise/documentation/gen_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
DEEP_BLUE  = RGBColor(0x1A, 0x37, 0x6C)   # #1A376C
MID_BLUE   = RGBColor(0x2E, 0x86, 0xC1)   # #2E86C1
ACCENT     = RGBColor(0x27, 0xAE, 0x60)   # #27AE60
DARK_GREY  = RGBColor(0x33, 0x33, 0x33)
MID_GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xD6, 0xEA, 0xF8)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic= italic
    if color:
        run.font.color.rgb = color

def heading(text, level=1, color=DEEP_BLUE, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    set_font(r, "Calibri", size or sizes.get(level, 12), bold=True, color=color)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    return p

def body(text, size=11, color=DARK_GREY, italic=False, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    set_font(r, "Calibri", size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, "Calibri", 11, color=DARK_GREY)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E86C1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def shade_cell(cell, hex_color="D6EAF8"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def simple_table(headers, rows, col_widths=None, header_bg="1A376C"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], header_bg)
        p  = hdr_cells[i].paragraphs[0]
        r  = p.add_run(h)
        set_font(r, "Calibri", 11, bold=True, color=WHITE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        bg = "EBF5FB" if ri % 2 == 0 else "FDFEFE"
        for ci, val in enumerate(cells):
            shade_cell(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(row[ci]) if ci < len(row) else "")
            set_font(r, "Calibri", 10.5, color=DARK_GREY)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n")

# UNESCO badge line
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("UNESCO GLOBAL MIL WEEK 2026 YOUTH CHALLENGE")
set_font(r2, "Calibri", 12, bold=True, color=MID_BLUE)
p2.paragraph_format.space_after = Pt(8)

# Project Title
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("FeedWise")
set_font(r3, "Calibri", 40, bold=True, color=DEEP_BLUE)
p3.paragraph_format.space_after = Pt(4)

# Tagline
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('"Don\'t Just Scroll. Think."')
set_font(r4, "Calibri", 16, italic=True, color=MID_BLUE)
p4.paragraph_format.space_after = Pt(16)

add_hr()

# Subtitle
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("A Practical Media & Information Literacy Platform for the AI Era")
set_font(r5, "Calibri", 14, color=DARK_GREY)
p5.paragraph_format.space_after = Pt(24)

# Team Section
p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run("Project Documentation")
set_font(r6, "Calibri", 13, bold=True, color=DEEP_BLUE)

p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
r7 = p7.add_run("Submitted for UNESCO Global MIL Week 2026 Youth Challenge")
set_font(r7, "Calibri", 12, color=MID_GREY)
p7.paragraph_format.space_after = Pt(20)

# Cover table — team info
cover_rows = [
    ["Team Leader",     "Upendra Singh Dhami"],
    ["Team Member",     "Bishal Regmi"],
    ["Team Member",     "Tulika Jha"],
    ["Team Member",     "Utsav Satgauwa Tharu"],
    ["Team Member",     "Pukar Pathak"],
    ["Challenge Track", "Games / Applications & Websites"],
    ["Submission Window","July 6 – August 16, 2026"],
    ["Category",        "Media & Information Literacy — AI Era"],
    ["Document Date",   "August 2026"],
]
simple_table(["Role", "Name / Detail"], cover_rows, col_widths=[1.8, 3.8])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  PAGE 2 — TEAM PROFILE
# ═══════════════════════════════════════════════════════════════════════════════
heading("Team Profile", level=1)
add_hr()

body(
    "FeedWise is developed by a team of five student developers and technologists united by a shared "
    "belief: that the most urgent skill for the next generation is not knowing how to use technology, "
    "but knowing how to think critically about the information that technology delivers."
)

heading("Team Members", level=2)

members = [
    ("Upendra Singh Dhami", "Team Leader & Lead Developer",
     "[Gmail: to be filled] | [WhatsApp: to be filled] | [College Email: to be filled]",
     "Upendra leads the FeedWise project — overseeing architecture, backend development, "
     "Flutter app development, and UNESCO submission strategy. Responsible for the full-stack "
     "design and final delivery of the platform."),
    ("Bishal Regmi", "Frontend Developer",
     "[Gmail: to be filled] | [WhatsApp: to be filled] | [College Email: to be filled]",
     "Bishal focuses on Flutter UI implementation, component library design, and "
     "cross-platform mobile and web experience optimization."),
    ("Tulika Jha", "UI/UX Designer & Content Strategist",
     "[Gmail: to be filled] | [WhatsApp: to be filled] | [College Email: to be filled]",
     "Tulika is responsible for user experience design, wireframing, interaction design, "
     "and MIL scenario content development."),
    ("Utsav Satgauwa Tharu", "Backend Developer",
     "[Gmail: to be filled] | [WhatsApp: to be filled] | [College Email: to be filled]",
     "Utsav contributes to the FastAPI backend, database integration, API endpoint "
     "development, and testing infrastructure."),
    ("Pukar Pathak", "AI & Data Engineer",
     "[Gmail: to be filled] | [WhatsApp: to be filled] | [College Email: to be filled]",
     "Pukar designs and implements the AI/NLP engine, prompt engineering, and "
     "data analytics pipelines that power FeedWise's intelligent features."),
]

for name, role, contact, desc in members:
    p = doc.add_paragraph()
    r = p.add_run(f"  {name} — {role}")
    set_font(r, "Calibri", 12, bold=True, color=DEEP_BLUE)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

    pc = doc.add_paragraph()
    rc = pc.add_run(f"  Contact: {contact}")
    set_font(rc, "Calibri", 10, italic=True, color=MID_GREY)
    pc.paragraph_format.space_after = Pt(2)

    pd = doc.add_paragraph()
    rd = pd.add_run(f"  {desc}")
    set_font(rd, "Calibri", 11, color=DARK_GREY)
    pd.paragraph_format.space_after = Pt(6)

body(
    "Note: Personal contact details (phone numbers, Gmail, and college email addresses) will be "
    "added by the team before final submission. Placeholder fields are marked above.",
    italic=True, color=MID_GREY
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  PAGE 3 — TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", level=1)
add_hr()

toc_items = [
    ("1", "Introduction", "4"),
    ("1.1", "What is FeedWise?", "4"),
    ("1.2", "Platform Vision", "4"),
    ("2", "Statement of Problem", "5"),
    ("2.1", "The Global MIL Crisis", "5"),
    ("2.2", "The AI Misinformation Challenge", "5"),
    ("2.3", "Gap in Existing Solutions", "6"),
    ("3", "Objectives", "6"),
    ("3.1", "Primary Objectives", "6"),
    ("3.2", "Secondary Objectives", "7"),
    ("3.3", "Success Metrics", "7"),
    ("4", "Target Audience", "8"),
    ("4.1", "Primary Users: Students & Youth", "8"),
    ("4.2", "Secondary Users: Teachers & Educators", "8"),
    ("4.3", "Tertiary Users: Administrators & MIL Professionals", "9"),
    ("4.4", "Audience Sizing & Inclusion Strategy", "9"),
    ("5", "Proposed Solution", "10"),
    ("5.1", "The FeedWise Approach", "10"),
    ("5.2", "Core Learning Loop", "10"),
    ("5.3", "Key Differentiators", "11"),
    ("6", "System Architecture", "11"),
    ("6.1", "High-Level Architecture", "11"),
    ("6.2", "Frontend Architecture", "12"),
    ("6.3", "Backend Architecture", "12"),
    ("6.4", "AI Engine Architecture", "13"),
    ("7", "Major Components", "13"),
    ("7.1", "Simulated Social Feed", "13"),
    ("7.2", "TrustLens Investigation Panel", "14"),
    ("7.3", "Decision & Consequence Engine", "14"),
    ("7.4", "MIL Academy", "14"),
    ("7.5", "Newsroom Mode", "15"),
    ("7.6", "Skill Radar & Progress System", "15"),
    ("7.7", "Badge & Achievement System", "15"),
    ("7.8", "Community Submission Module", "16"),
    ("8", "User Roles", "16"),
    ("8.1", "General User (Student)", "16"),
    ("8.2", "Teacher / Educator", "17"),
    ("8.3", "Administrator", "17"),
    ("8.4", "Reviewer / MIL Professional", "18"),
    ("8.5", "Role-Based Access Control", "18"),
    ("9", "Relational Database Schema", "18"),
    ("9.1", "Entity Relationship Overview", "19"),
    ("9.2", "Core Tables", "19"),
    ("9.3", "Learning & Progress Tables", "20"),
    ("9.4", "Classroom & Analytics Tables", "20"),
    ("10", "Technology Stack", "21"),
    ("11", "Security Architecture", "21"),
    ("11.1", "Authentication & Authorization", "21"),
    ("11.2", "Row Level Security (RLS)", "22"),
    ("11.3", "OWASP Compliance", "22"),
    ("12", "Business Perspective", "22"),
    ("12.1", "Revenue Model", "22"),
    ("12.2", "Market Opportunity", "23"),
    ("12.3", "Sustainability Plan", "23"),
    ("12.4", "Partnership Strategy", "24"),
    ("13", "Creativity & Innovation", "24"),
    ("13.1", "Creative Design Principles", "24"),
    ("13.2", "Key Innovations", "25"),
    ("13.3", "FeedWise vs. Existing Landscape", "25"),
    ("14", "Feasibility Assessment", "26"),
    ("14.1", "Technical Feasibility", "26"),
    ("14.2", "Financial & Operational Feasibility", "26"),
    ("14.3", "Market Feasibility", "27"),
    ("14.4", "Risk Assessment & Mitigation", "27"),
    ("15", "Future Development", "27"),
    ("15.1", "Phase Roadmap", "27"),
    ("15.2", "Planned Features", "28"),
    ("15.3", "Research & Impact Measurement", "28"),
    ("16", "Conclusion", "29"),
    ("17", "References", "30"),
    ("—", "Appendix A: Screenshots", "31"),
]

for num, title, pg in toc_items:
    p = doc.add_paragraph()
    is_main = not "." in num and num != "—"
    r = p.add_run(f"{num}   {title}")
    set_font(r, "Calibri", 11, bold=is_main, color=DEEP_BLUE if is_main else DARK_GREY)

    # Tab + page number
    r2 = p.add_run(f"{'.' * max(1, 55 - len(num) - len(title))}  {pg}")
    set_font(r2, "Calibri", 11, color=MID_GREY)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. Introduction", level=1)
add_hr()

heading("1.1 What is FeedWise?", level=2)
body(
    "FeedWise is an interactive, AI-assisted Media and Information Literacy (MIL) platform that "
    "simulates the modern digital information environment for young learners. It is submitted as "
    "an entry to the UNESCO Global MIL Week 2026 Youth Challenge under the track "
    "'Games / Applications & Websites'."
)
body(
    "At its core, FeedWise is a social media simulation — a safe, controlled, consequence-modelled "
    "digital space where students encounter realistic information scenarios, investigate them using "
    "structured evidence tools, make decisions, experience consequences, and immediately receive "
    "targeted MIL lessons. The platform does not tell users what is true or false. It teaches them "
    "how to evaluate evidence, recognize manipulation, and make informed judgements — skills that "
    "are indispensable in the AI era."
)

heading("1.2 Platform Vision", level=2)
body(
    "FeedWise envisions a world where every young person possesses the cognitive tools to navigate "
    "the information ecosystem responsibly. The platform is designed to grow from an MVP learning "
    "tool into a global MIL ecosystem — multilingual, community-powered, classroom-integrated, "
    "and professionally curated by educators, journalists, and MIL researchers."
)
body(
    "FeedWise is not just an app. It is a methodology: evidence-first, consequence-aware, "
    "AI-assisted and human-reviewed. This methodology ensures that every interaction on the "
    "platform is an authentic learning experience, not a gamified quiz."
)

body(
    "The platform is built for three interconnected communities: students and youth who experience "
    "the simulated feed; teachers and educators who monitor, guide, and assign targeted learning; "
    "and administrators and MIL professionals who create, review, and maintain the quality of every "
    "scenario on the platform."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — STATEMENT OF PROBLEM
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. Statement of Problem", level=1)
add_hr()

heading("2.1 The Global MIL Crisis", level=2)
body(
    "We are living through the most complex information environment in human history. Young people "
    "aged 14–25 spend an average of 7 hours per day consuming digital content — social media feeds, "
    "messaging apps, video platforms, and online news. Within this vast information stream, "
    "misinformation, disinformation, and AI-generated synthetic content circulate at unprecedented "
    "speed and volume."
)
body(
    "According to UNESCO's 2023 MIL report, fewer than 20% of young people in developing nations "
    "have received any formal media literacy education. Yet these same young people are primary "
    "consumers and amplifiers of digital content — often sharing information before verifying it, "
    "rarely questioning the source, and increasingly unable to distinguish AI-generated content "
    "from authentic human expression."
)
body(
    "The consequences are severe: public health crises amplified by medical misinformation; "
    "democratic processes undermined by coordinated disinformation campaigns; social trust eroded "
    "by deep fake videos; and entire communities polarized by algorithmically engineered outrage."
)

heading("2.2 The AI Misinformation Challenge", level=2)
body(
    "The rapid democratization of generative AI has introduced a new dimension to the "
    "misinformation problem. Large language models can produce convincing news articles, "
    "authoritative-sounding academic citations, and plausible social media posts at zero marginal "
    "cost. AI image generators produce photorealistic images of events that never happened. "
    "Voice cloning tools reproduce public figures saying things they never said."
)
body(
    "The result is an 'epistemic crisis' — a situation where the average person can no longer "
    "rely on intuitive signals (Does this look real? Does this sound right?) to evaluate "
    "information. The skills required to navigate this environment are fundamentally different "
    "from traditional media literacy and must be explicitly taught and practiced."
)

heading("2.3 Gap in Existing Solutions", level=2)
body(
    "While media literacy education has been growing, existing solutions fall short in critical ways:"
)
bullet("Textbook-based MIL curricula teach theory but provide no practice environment.")
bullet("Fact-checking apps give binary true/false verdicts — they do the thinking for users rather than teaching them to think.")
bullet("Existing educational games are either too simplistic or fail to connect gameplay to genuine MIL skill development.")
bullet("Few platforms integrate teacher oversight with student learning analytics.")
bullet("Almost no solution addresses AI-generated content, deepfakes, and synthetic media as a specific learning domain.")
bullet("Classroom integration is rare — most tools are designed for individual use, not structured educational settings.")

body(
    "FeedWise exists to fill this gap. It is the first platform to combine a realistic social "
    "media simulation, a structured evidence investigation framework, a consequence engine, "
    "professional MIL educator review, real-time AI assistance, and teacher classroom integration "
    "into a single, coherent learning system."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — OBJECTIVES
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. Objectives", level=1)
add_hr()

heading("3.1 Primary Objectives", level=2)
bullet("Provide young people aged 14–25 with an immersive, practice-based MIL learning experience through a simulated social media environment.")
bullet("Develop measurable competency in five MIL skill dimensions: Source Verification, Evidence Evaluation, AI Literacy, Bias Detection, and Digital Safety.")
bullet("Enable teachers and educators to integrate FeedWise into classroom instruction with real-time student monitoring, assignment tools, and skill gap analysis.")
bullet("Build an evidence-first AI engine that assists — but never replaces — human analysis, ensuring users learn to think rather than rely on automated verdicts.")
bullet("Create a sustainable, open, community-powered MIL scenario ecosystem that grows with contributions from educators, journalists, and MIL professionals worldwide.")
bullet("Deliver a platform that works across devices and languages, prioritizing accessibility and inclusion for learners in under-resourced settings.")

heading("3.2 Secondary Objectives", level=2)
bullet("Establish a replicable methodology for consequence-aware MIL education that can be adapted by schools, NGOs, and governments.")
bullet("Generate research-quality data on student MIL skill development patterns and learning outcomes.")
bullet("Build a badge and achievement system that motivates continued engagement through intrinsic recognition of skill mastery.")
bullet("Provide a Newsroom Mode that builds empathy for the challenges of responsible journalism and editorial decision-making.")

heading("3.3 Success Metrics", level=2)
simple_table(
    ["Metric", "Target (Year 1)", "Measurement Method"],
    [
        ["Active users", "10,000+", "Platform analytics"],
        ["Avg. MIL skill improvement", "25%+ across 5 dimensions", "Pre/post baseline test"],
        ["Teacher classrooms integrated", "500+", "Account registrations"],
        ["Scenarios published", "200+", "Content database"],
        ["Daily active engagement", "15+ minutes/session", "Session analytics"],
        ["Community submissions", "1,000+", "Submission table"],
    ],
    col_widths=[2.2, 1.8, 2.6]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — TARGET AUDIENCE
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. Target Audience", level=1)
add_hr()

body(
    "FeedWise is designed for a diverse, interconnected set of users who together form a "
    "complete Media and Information Literacy learning ecosystem. Understanding the distinct "
    "needs, contexts, and goals of each audience group drives every design and feature decision "
    "on the platform."
)

heading("4.1 Primary Users: Students & Youth (Age 14–25)", level=2)
body(
    "Young people aged 14–25 are the primary and most critical audience for FeedWise. This "
    "demographic is the heaviest consumer of social media content globally, the most frequent "
    "sharer of unverified information, and simultaneously the most receptive to new ways of "
    "learning when presented in an engaging, game-like format."
)
simple_table(
    ["Characteristic", "Detail"],
    [
        ["Age Range",               "14–25 years (core); 12–13 with teacher supervision"],
        ["Device Usage",            "Primarily mobile (Android-first); web-accessible for all platforms"],
        ["Learning Context",        "Individual use, classroom, after-school programs, community groups"],
        ["Language",                "English (launch); Nepali, Hindi, Bengali, Arabic (Phase 2 roadmap)"],
        ["Geography — Primary",     "Nepal, India, Bangladesh — high digital adoption, low MIL infrastructure"],
        ["Geography — Secondary",   "Sub-Saharan Africa, Southeast Asia, Eastern Europe, Latin America"],
        ["Connectivity",            "Low-bandwidth-friendly design; offline caching in development roadmap"],
        ["Education Level",         "Secondary school (Grade 8+) through undergraduate university"],
        ["Motivation Driver",       "Skill growth, badges, leaderboard, classroom credit, career readiness"],
    ],
    col_widths=[2.4, 4.2]
)
body(
    "FeedWise meets youth users where they already are — on social media-style interfaces — "
    "using familiar patterns (feeds, likes, shares, badges) to build essential critical thinking skills."
)

heading("4.2 Secondary Users: Teachers & Educators", level=2)
body(
    "Teachers use FeedWise Studio to monitor student progress, identify skill gaps, assign "
    "targeted challenges, and integrate MIL learning into formal curricula."
)
simple_table(
    ["Characteristic", "Detail"],
    [
        ["Profile",          "Secondary school & university teachers; media studies; social studies; ICT educators; NGO facilitators"],
        ["Primary Need",     "Class monitoring, skill gap analysis, assignment tools, exportable progress reports"],
        ["Technical Comfort","Moderate — comfortable with Google Classroom or LMS; standard web browser"],
        ["Core Pain Solved", "No existing tool gives teachers real-time MIL skill analytics for their students"],
        ["Geography",        "Wherever FeedWise is deployed; initial pilot targeting Nepali and Indian schools"],
    ],
    col_widths=[2.2, 4.4]
)

heading("4.3 Tertiary Users: Administrators & MIL Professionals", level=2)
body(
    "Administrators and MIL professionals ensure the quality, accuracy, and educational integrity "
    "of every scenario. They are typically experienced journalists, UNESCO-trained MIL educators, "
    "academic researchers in information science, or content moderation specialists."
)
simple_table(
    ["Characteristic", "Detail"],
    [
        ["Profile",          "MIL researchers, journalists, UNESCO educators, digital rights advocates"],
        ["Primary Need",     "Scenario CMS, review pipeline, content quality control, AI output oversight"],
        ["Key Value",        "Platform credibility depends on their review — the human check on AI"],
        ["Engagement",       "Salaried partner institutions or volunteer with professional development recognition"],
        ["Reach",            "Global — remote content review possible from any location"],
    ],
    col_widths=[2.2, 4.4]
)

heading("4.4 Audience Sizing & Inclusion Strategy", level=2)
body(
    "FeedWise's core MIL learning experience is free, free forever, and designed to work on "
    "entry-level Android devices with limited connectivity. Inclusion is not an afterthought — "
    "it is the product mission.",
    bold=True, color=DEEP_BLUE
)
simple_table(
    ["Audience Segment",           "Global Scale",              "Accessibility Strategy"],
    [
        ["Youth 14–25",                "1.8 billion globally",     "Free access forever; mobile-first; low-bandwidth design"],
        ["Teachers / Educators",       "90 million globally",      "Free teacher accounts; institutional onboarding support"],
        ["MIL Professionals",          "500K+ (UNESCO estimate)",  "Volunteer reviewer program; professional development credit"],
        ["Marginalized Communities",   "Billions in low-resource settings", "Offline mode roadmap; multilingual content"],
        ["Students with Disabilities", "~15% of all learners",    "Screen reader support; high-contrast mode; text scaling"],
    ],
    col_widths=[2.2, 1.8, 2.6]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — PROPOSED SOLUTION
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. Proposed Solution", level=1)
add_hr()

heading("5.1 The FeedWise Approach", level=2)
body(
    "FeedWise proposes a simulation-first, evidence-based approach to MIL education. Rather than "
    "delivering content about media literacy, FeedWise puts students inside a realistic information "
    "environment where they must practice MIL skills to navigate it successfully."
)
body(
    "The platform is grounded in three core principles:"
)
bullet("Evidence-First: Every decision must be supported by evidence. The platform rewards investigation, not guessing.")
bullet("Consequence-Aware: Every decision produces measurable simulated impact, making the cost of poor information hygiene tangible and immediate.")
bullet("Human-Reviewed: Every scenario on the platform is reviewed and approved by qualified MIL educators before publication, ensuring educational integrity.")

heading("5.2 Core Learning Loop", level=2)
body(
    "Every user session on FeedWise follows a carefully designed eight-step learning loop:"
)

loop_rows = [
    ["1", "DISCOVER",     "User sees a realistic simulated social media post in their feed"],
    ["2", "INTERACT",     "User reads the post, notices engagement numbers, trending tags"],
    ["3", "INVESTIGATE",  "User opens TrustLens — examines source, author, date, evidence, language signals"],
    ["4", "DECIDE",       "User selects: SHARE / VERIFY / REPORT / IGNORE"],
    ["5", "CONSEQUENCE",  "System reveals impact: reach numbers, shares, credibility delta, community effects"],
    ["6", "LEARN",        "Targeted MIL lesson delivered, tied specifically to this scenario's key skill"],
    ["7", "IMPROVE",      "5-dimension skill radar updates to reflect performance"],
    ["8", "NEXT",         "User is served next scenario, difficulty calibrated to current skill level"],
]
simple_table(["Step", "Stage", "What Happens"], loop_rows, col_widths=[0.5, 1.4, 4.7])

heading("5.3 Key Differentiators", level=2)
simple_table(
    ["What Exists Today", "What FeedWise Does Differently"],
    [
        ["Fact-checking apps answer 'Is this true?'", "FeedWise asks 'Can YOU investigate it?'"],
        ["Educational content: 'Here are the rules'", "FeedWise: 'Try it yourself — see consequences'"],
        ["Binary true/false verdicts", "Evidence-based assessment with nuanced signals"],
        ["AI tools replace human thinking", "AI assists investigation; human MIL experts review all content"],
        ["Individual learning apps", "Full classroom integration with teacher analytics"],
        ["Entertainment-focused games", "Every game mechanic tied to measurable MIL skill"],
        ["English-only platforms", "Multi-language architecture (English, Nepali, and expanding)"],
    ],
    col_widths=[2.8, 3.8]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading("6. System Architecture", level=1)
add_hr()

heading("6.1 High-Level Architecture", level=2)
body(
    "FeedWise is a three-tier, cloud-native architecture consisting of two Flutter client "
    "applications, one FastAPI backend server, and a Supabase-managed PostgreSQL database. "
    "All communication between clients and backend occurs over HTTPS REST APIs."
)

# ASCII art approximation in docx
arch_text = (
    "FEEDWISE — HIGH-LEVEL ARCHITECTURE\n\n"
    "┌─────────────────────┐    ┌─────────────────────┐\n"
    "│  USER APPLICATION   │    │  ADMIN / TEACHER     │\n"
    "│  Flutter (Android + │    │  STUDIO (Flutter Web)│\n"
    "│  Web + iOS)         │    │                      │\n"
    "│  ─ Feed & Invest.   │    │  ─ Scenario CMS      │\n"
    "│  ─ Academy          │    │  ─ Teacher Dashboard │\n"
    "│  ─ Profile & Badges │    │  ─ Analytics         │\n"
    "└──────────┬──────────┘    └──────────┬───────────┘\n"
    "           │  HTTPS / REST API         │\n"
    "           └─────────┬─────────────────┘\n"
    "                     │\n"
    "          ┌──────────▼──────────┐\n"
    "          │    FASTAPI BACKEND  │\n"
    "          │  ┌───┐  ┌───┐  ┌───┐  │\n"
    "          │  │SCE│  │AI │  │EVD│  │\n"
    "          │  │NAR│  │ENG│  │ENG│  │\n"
    "          │  │IO │  │INE│  │INE│  │\n"
    "          │  └───┘  └───┘  └───┘  │\n"
    "          └──────────┬──────────┘\n"
    "                     │\n"
    "          ┌──────────▼──────────┐\n"
    "          │  SUPABASE           │\n"
    "          │  PostgreSQL + Auth  │\n"
    "          │  + Storage + RLS    │\n"
    "          └─────────────────────┘"
)

p_arch = doc.add_paragraph()
r_arch = p_arch.add_run(arch_text)
set_font(r_arch, "Courier New", 9, color=DARK_GREY)
p_arch.paragraph_format.space_after = Pt(8)

heading("6.2 Frontend Architecture", level=2)
body(
    "Both Flutter applications use the Riverpod state management library for reactive, "
    "dependency-injected state, and GoRouter for declarative, type-safe navigation. "
    "The shared design system implements a consistent dark-mode-first design language "
    "with a custom color palette, Outfit typography, and a comprehensive widget library."
)
simple_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["State Management", "Riverpod + AsyncNotifier", "Reactive, testable state"],
        ["Navigation", "GoRouter", "Declarative routing with guards"],
        ["HTTP Client", "Dio + Retrofit", "Type-safe API calls"],
        ["Local Storage", "SharedPreferences + Sqflite", "Offline data caching"],
        ["Authentication", "Supabase Auth SDK", "JWT-based auth"],
        ["Animations", "Rive + flutter_animate", "Micro-interactions"],
        ["Charts", "fl_chart", "Skill radar, analytics"],
    ],
    col_widths=[1.8, 2.0, 2.8]
)

heading("6.3 Backend Architecture", level=2)
body(
    "The FastAPI backend uses a strict layered architecture: API routes call service layer only; "
    "the service layer orchestrates repositories and the AI engine; repositories handle all "
    "database operations. No business logic exists in route handlers or database queries."
)
simple_table(
    ["Layer", "Responsibility"],
    [
        ["API Routes (/api/v1/)", "HTTP request handling, input validation, response formatting"],
        ["Service Layer", "Business logic, orchestration, AI engine calls"],
        ["Repository Layer", "Database queries via Supabase client — no business logic"],
        ["AI Engine", "LLM API calls, text analysis, explanation generation"],
        ["Middleware", "Auth verification, CORS, rate limiting, request logging"],
    ],
    col_widths=[2.5, 4.1]
)

heading("6.4 AI Engine Architecture", level=2)
body(
    "FeedWise's AI engine operates on a fundamental principle: evidence-first, AI-assisted, "
    "human-reviewed. The AI engine does NOT issue true/false verdicts — doing so would "
    "undermine FeedWise's core educational mission of teaching critical thinking, not replacing it."
)
body("The AI engine provides four capabilities:")
bullet("Text Signal Analysis: Analyzes post language for emotional framing, certainty level, urgency signals, and clickbait patterns.")
bullet("Claim Extraction: Identifies the primary factual claim within a social media post to guide the investigation panel.")
bullet("Explanation Generation: Drafts educational MIL lesson text for admin review and approval.")
bullet("Translation Assistance: Provides draft translation of scenarios into additional languages, reviewed before publication.")
body(
    "All AI responses are validated against JSON schemas. Any AI failure falls back gracefully "
    "to pre-authored content, ensuring the platform functions without AI availability.",
    italic=False
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — MAJOR COMPONENTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("7. Major Components", level=1)
add_hr()

heading("7.1 Simulated Social Feed", level=2)
body(
    "The social feed is the primary interface of FeedWise. It renders realistic social media "
    "post cards with author names, headlines, body text, engagement numbers (likes, shares, "
    "comments), publication dates, source indicators, trending badges, and category tags. "
    "Posts are carefully designed to mimic the visual language of real social platforms "
    "while remaining fictional and educational. The feed uses intelligent scenario serving: "
    "each user sees scenarios calibrated to their current skill level, avoiding repetition "
    "and ensuring progressive challenge."
)

heading("7.2 TrustLens Investigation Panel", level=2)
body(
    "TrustLens is FeedWise's signature investigation tool — a structured, multi-section panel "
    "that guides users through a professional information verification workflow. It exposes "
    "five investigation dimensions:"
)
bullet("Source Transparency: Who published this? Is the source identified? Is contact information available?")
bullet("Author Credibility: Is the author known? Do they have a verifiable track record?")
bullet("Date Verification: Can the publication date be independently confirmed? Is this old content being recirculated?")
bullet("Evidence Assessment: What primary and secondary evidence exists? Does the evidence support or contradict the claim?")
bullet("Language Signals: Does the text use emotional language, absolute certainty, urgency, or clickbait patterns?")
body(
    "The AI engine assists TrustLens by analyzing post language in real time. However, all "
    "signals are presented as evidence for the user to evaluate — never as a final verdict."
)

heading("7.3 Decision & Consequence Engine", level=2)
body(
    "After investigating, the user selects one of four decisions: SHARE, VERIFY, REPORT, or IGNORE. "
    "Each decision triggers the Consequence Engine, which renders a consequence card showing:"
)
bullet("Information Reach: How many simulated people saw this post after the user's action.")
bullet("Share Spread: How many times the content was reshared as a result.")
bullet("Credibility Delta: How the user's credibility score changes based on decision quality.")
bullet("Impact Text: A narrative description of the real-world effect this decision would have.")
bullet("Missed Clues: Key evidence signals the user overlooked, explained clearly.")
body(
    "This immediate feedback loop makes the cost of poor information hygiene visceral and memorable — "
    "far more effective than abstract warnings about the dangers of misinformation."
)

heading("7.4 MIL Academy", level=2)
body(
    "The MIL Academy is the structured educational layer of FeedWise. It organizes learning "
    "content into modules aligned with the five MIL skill dimensions: Source Verification, "
    "Evidence Evaluation, AI Literacy, Bias Detection, and Digital Safety. Each module contains "
    "bite-sized lessons (3–5 minutes each), scenario challenges, and quizzes. Academy content "
    "is authored by MIL professionals and reviewed for accuracy, cultural appropriateness, "
    "and age-suitability. The Academy tracks completion and connects lesson outcomes to the "
    "user's skill radar, ensuring every learning activity has measurable impact."
)

heading("7.5 Newsroom Mode", level=2)
body(
    "Newsroom Mode ('Newsroom Zero') places the user in the role of a news editor deciding "
    "which stories to publish. This mode builds editorial empathy — users experience the "
    "pressures, trade-offs, and ethical dilemmas that real journalists face. Stories contain "
    "conflicting evidence, time pressure signals, and competing public interest considerations. "
    "Decisions in Newsroom Mode affect simulated publication metrics, readership trust scores, "
    "and advertiser relationships, creating a holistic understanding of the media industry ecosystem."
)

heading("7.6 Skill Radar & Progress System", level=2)
body(
    "Every action on FeedWise contributes to a five-dimensional skill radar that visualizes "
    "the user's MIL competency profile. The five axes are: Source Verification, Evidence "
    "Evaluation, AI Literacy, Bias Detection, and Digital Safety. Skill scores are computed "
    "from decision accuracy, investigation depth, lesson completion, and quiz performance. "
    "XP and level progression provide additional motivation, while streak tracking rewards "
    "consistent engagement. The skill radar is visible to both students (for self-reflection) "
    "and teachers (for monitoring class-wide skill gaps)."
)

heading("7.7 Badge & Achievement System", level=2)
body(
    "FeedWise includes a comprehensive badge system that recognizes skill mastery and learning "
    "milestones. Badges are organized across categories:"
)
simple_table(
    ["Badge", "Description", "Trigger"],
    [
        ["Source Detective", "Master of source verification", "95%+ accuracy on source scenarios"],
        ["AI Aware", "Identifies AI-generated content", "Complete AI Literacy module"],
        ["Fact Finder", "Evidence evaluation expert", "25+ scenarios verified correctly"],
        ["Bias Buster", "Recognizes framing & manipulation", "10 bias scenarios completed"],
        ["Community Contributor", "Helps others learn", "5 community submissions approved"],
        ["Newsroom Pro", "Editorial decision-making expert", "Complete 10 newsroom scenarios"],
        ["Streak Champion", "Consistent daily learner", "30-day login streak"],
    ],
    col_widths=[1.8, 2.6, 2.2]
)

heading("7.8 Community Submission Module", level=2)
body(
    "Advanced users can submit real-world information claims for community investigation. "
    "Submissions go through a moderation pipeline: submitted → reviewed by MIL professionals → "
    "approved for community investigation → published as community challenges. This creates a "
    "living, user-generated library of contemporary MIL scenarios, ensures the platform remains "
    "current, and gives learners ownership over the educational ecosystem."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — USER ROLES
# ═══════════════════════════════════════════════════════════════════════════════
heading("8. User Roles", level=1)
add_hr()

body(
    "FeedWise implements a five-tier role system managed through Role-Based Access Control "
    "(RBAC). Each role defines what a user can see, create, modify, and access. Roles are "
    "stored in the user_roles table and enforced at both the API layer and via database-level "
    "Row Level Security (RLS) policies."
)

heading("8.1 General User (Student / Youth)", level=2)
body(
    "The general user role is the primary learner — a student, young person, or any member "
    "of the public who engages with FeedWise as a media literacy learning tool."
)
simple_table(
    ["Capability", "Details"],
    [
        ["Feed Access", "View and interact with published scenarios in the social feed"],
        ["Investigation", "Use TrustLens to investigate any feed scenario"],
        ["Decision Making", "Submit decisions (Share/Verify/Report/Ignore) on scenarios"],
        ["Academy Learning", "Access and complete all MIL Academy modules and lessons"],
        ["Newsroom Mode", "Participate in Newsroom Zero editorial simulations"],
        ["Progress Tracking", "View personal skill radar, XP, level, and badge collection"],
        ["Community Mode", "Submit real-world claims for community review (after Level 5)"],
        ["Profile Management", "Edit personal profile, goals, language preferences"],
        ["Notifications", "Receive learning reminders, badge awards, assignment notifications"],
    ],
    col_widths=[2.2, 4.4]
)

heading("8.2 Teacher / Educator Role", level=2)
body(
    "Teachers use the FeedWise Studio (the admin/teacher portal) to manage classrooms, "
    "monitor student progress, identify skill gaps, and assign targeted learning challenges."
)
simple_table(
    ["Capability", "Details"],
    [
        ["Class Management", "Create and manage classes; add/remove students via invite code"],
        ["Student Monitoring", "View individual student skill radars, progress, and decision history"],
        ["Skill Gap Analysis", "See class-wide skill dimension averages; identify weakest areas"],
        ["Assignment Creation", "Assign specific scenarios or Academy modules to the whole class"],
        ["Analytics Dashboard", "Track engagement metrics, completion rates, and improvement trends"],
        ["Lesson Recommendations", "Receive AI-generated recommendations for targeted interventions"],
        ["Scenario Preview", "Preview any published scenario before assigning it to students"],
        ["Report Generation", "Export student progress reports for institutional records"],
    ],
    col_widths=[2.2, 4.4]
)

heading("8.3 Administrator Role", level=2)
body(
    "Administrators are the platform stewards — typically MIL professionals, journalists, "
    "or senior educators. They maintain the quality, accuracy, and integrity of every piece "
    "of content on FeedWise."
)
simple_table(
    ["Capability", "Details"],
    [
        ["Scenario CMS", "Create, edit, delete, and publish scenarios via full CRUD interface"],
        ["Review Pipeline", "Manage scenario lifecycle: Draft → Reviewed → Approved → Published"],
        ["Evidence Management", "Attach, edit, and verify evidence items for each scenario"],
        ["AI Engine Oversight", "Review and approve AI-generated draft content before publication"],
        ["User Management", "View, suspend, or ban users; manage role assignments"],
        ["Community Moderation", "Review and approve or reject community-submitted scenarios"],
        ["Analytics Access", "Platform-wide analytics: engagement, learning outcomes, content performance"],
        ["Audit Log", "Full audit trail of all administrative actions for accountability"],
        ["Localization Management", "Manage translated scenario content across all supported languages"],
        ["Platform Settings", "Configure feature flags, difficulty calibration, and moderation thresholds"],
    ],
    col_widths=[2.4, 4.2]
)

heading("8.4 Reviewer / MIL Professional Role", level=2)
body(
    "Reviewers are specialized users — typically academic MIL researchers, experienced journalists, "
    "or certified educators — who focus specifically on content quality review. They can approve "
    "or reject scenarios in the review pipeline, annotate evidence items, and flag content for "
    "additional verification. They do not have full admin access (user management, platform settings)."
)

heading("8.5 Role-Based Access Control (RBAC)", level=2)
simple_table(
    ["Permission", "Student", "Teacher", "Reviewer", "Admin"],
    [
        ["View Feed & Scenarios",     "Yes", "Yes", "Yes", "Yes"],
        ["Make Decisions on Scenarios","Yes", "Yes", "Yes", "Yes"],
        ["Access Academy",             "Yes", "Yes", "Yes", "Yes"],
        ["View Own Analytics",         "Yes", "Yes", "Yes", "Yes"],
        ["View Class Analytics",       "No",  "Yes", "No",  "Yes"],
        ["Create/Edit Scenarios",      "No",  "No",  "No",  "Yes"],
        ["Review Scenario Pipeline",   "No",  "No",  "Yes", "Yes"],
        ["Manage Users",               "No",  "No",  "No",  "Yes"],
        ["Platform-Wide Analytics",    "No",  "No",  "No",  "Yes"],
        ["Community Submissions",      "Yes*","No",  "Yes", "Yes"],
    ],
    col_widths=[2.8, 0.9, 0.9, 1.0, 0.9]
)
body("* Community submissions enabled after student reaches Level 5.", italic=True, color=MID_GREY)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════════════════════════
heading("9. Relational Database Schema", level=1)
add_hr()

heading("9.1 Entity Relationship Overview", level=2)
body(
    "FeedWise uses PostgreSQL hosted on Supabase, with Row Level Security (RLS) enforced "
    "at the database level. The schema is organized into six groups of tables: Core (users "
    "and roles), Content (scenarios and evidence), Learning (decisions and progress), "
    "Academy (modules and lessons), Classroom (classes and assignments), and Analytics."
)

er_text = (
    "ER DIAGRAM — SIMPLIFIED\n\n"
    "auth.users (Supabase)\n"
    "     │  1:1\n"
    "  profiles ──── user_roles\n"
    "     │\n"
    "     ├── decisions ──────────── scenarios\n"
    "     │                              │\n"
    "     ├── skill_progress             ├── content_items\n"
    "     │                              ├── evidence\n"
    "     ├── user_progress              ├── consequences\n"
    "     │                              └── lessons\n"
    "     ├── user_badges\n"
    "     │\n"
    "     ├── class_members ──────── classes ── assignments\n"
    "     │\n"
    "     └── community_submissions\n"
)
p_er = doc.add_paragraph()
r_er = p_er.add_run(er_text)
set_font(r_er, "Courier New", 9, color=DARK_GREY)
p_er.paragraph_format.space_after = Pt(8)

heading("9.2 Core Tables — Users & Roles", level=2)
simple_table(
    ["Table", "Key Columns", "Purpose"],
    [
        ["profiles", "id (FK → auth.users), username, display_name, age_group, locale, avatar_url, goals[], onboarding_completed", "Extended user profile data"],
        ["user_roles", "id, user_id (FK), role (student|teacher|reviewer|admin), granted_by, granted_at", "Multi-role assignments per user"],
    ],
    col_widths=[1.5, 3.8, 1.5]
)

heading("Content Tables — Scenarios & Evidence", level=2)
simple_table(
    ["Table", "Key Columns", "Purpose"],
    [
        ["scenarios", "id, title, category, difficulty, language, skill_tags[], claim_text, source_transparency, assessment, correct_decision, status (draft→published)", "Core scenario definitions"],
        ["content_items", "id, scenario_id (FK), author_name, headline, body, image_url, source_url, likes_count, shares_count, is_trending", "Simulated social media posts"],
        ["evidence", "id, scenario_id (FK), category, status, label, value, explanation", "Evidence items per scenario"],
        ["consequences", "id, scenario_id (FK), decision_type, reach_count, shares_count, credibility_delta, impact_text, missed_clues[]", "Consequence data per decision"],
        ["lessons", "id, scenario_id (FK), skill, title, tips[], key_takeaway, quiz_data", "MIL lessons per scenario"],
    ],
    col_widths=[1.5, 3.6, 1.7]
)

heading("9.3 Learning & Progress Tables", level=2)
simple_table(
    ["Table", "Key Columns", "Purpose"],
    [
        ["decisions", "id, user_id, scenario_id, decision (share|verify|report|ignore), is_correct, score, time_spent_sec, created_at", "Record of every user decision"],
        ["skill_progress", "id, user_id, skill (5 dimensions), score (0–100), updated_at", "Per-skill scores for skill radar"],
        ["user_progress", "id, user_id, xp_total, level, streak_days, last_active", "XP, level, and streak tracking"],
        ["user_badges", "id, user_id, badge_id, earned_at", "Badges earned by user"],
        ["badge_definitions", "id, name, description, icon, trigger_type, trigger_value", "All possible badges"],
        ["user_lesson_completions", "id, user_id, lesson_id, score, completed_at", "Academy lesson completion log"],
    ],
    col_widths=[2.0, 3.4, 1.4]
)

heading("9.4 Classroom & Analytics Tables", level=2)
simple_table(
    ["Table", "Key Columns", "Purpose"],
    [
        ["classes", "id, teacher_id, name, invite_code, created_at", "Teacher-managed classes"],
        ["class_members", "id, class_id, student_id, joined_at", "Students enrolled in class"],
        ["assignments", "id, class_id, teacher_id, scenario_id|module_id, due_date, created_at", "Teacher assignments to classes"],
        ["community_submissions", "id, submitter_id, claim_text, source_url, status (submitted→approved), reviewed_by", "User-submitted scenarios"],
        ["analytics_events", "id, user_id, event_type, event_data (JSONB), created_at", "Raw event log for analytics"],
        ["audit_log", "id, actor_id, action_type, target_table, target_id, old_data, new_data, created_at", "Admin action audit trail"],
        ["notifications", "id, user_id, type, title, body, is_read, created_at", "User notification queue"],
    ],
    col_widths=[2.1, 3.4, 1.3]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════════
heading("10. Technology Stack", level=1)
add_hr()

simple_table(
    ["Layer", "Technology", "Version", "Reason"],
    [
        ["User App",          "Flutter (Dart)",        "3.27+",   "Cross-platform: iOS, Android, Web from one codebase"],
        ["Admin Studio",      "Flutter Web",           "3.27+",   "Consistent UI system with user app"],
        ["State Management",  "Riverpod",              "2.5+",    "Reactive, testable, dependency-injected state"],
        ["Navigation",        "GoRouter",              "13+",     "Declarative routing with role-based guards"],
        ["Backend API",       "FastAPI (Python)",      "0.115+",  "Async, auto-docs, Pydantic-native, fast"],
        ["Validation",        "Pydantic v2",           "2.11+",   "Type-safe data validation"],
        ["Database",          "PostgreSQL (Supabase)", "15+",     "Managed, row-level security, real-time"],
        ["Authentication",    "Supabase Auth",         "2.x",     "JWT, OAuth2, Google Sign-In"],
        ["File Storage",      "Supabase Storage",      "—",       "Scenario images, user avatars"],
        ["AI/NLP",            "OpenAI / Gemini API",   "GPT-4o / Gemini 2.0", "Swappable LLM client"],
        ["Container",         "Docker",                "latest",  "Reproducible backend deployment"],
        ["Hosting (API)",     "Render.com",            "—",       "Free tier, auto-deploy from Git"],
        ["Hosting (Web)",     "Firebase Hosting",      "—",       "Global CDN, free tier"],
        ["Testing",           "pytest + httpx",        "—",       "Async-compatible backend testing"],
        ["CI/CD",             "GitHub Actions",        "—",       "Automated test + deploy pipeline"],
    ],
    col_widths=[1.8, 1.8, 1.4, 2.6]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — SECURITY ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading("11. Security Architecture", level=1)
add_hr()

heading("11.1 Authentication & Authorization", level=2)
body(
    "All authentication is handled by Supabase Auth using JWT tokens. The FastAPI backend "
    "verifies every JWT using the Supabase JWT secret on every protected request. Role claims "
    "are extracted from the verified JWT and checked against required role permissions before "
    "any service logic executes. Tokens expire after configurable intervals and are refreshed "
    "client-side."
)

heading("11.2 Row Level Security (RLS)", level=2)
body(
    "Row Level Security is enabled on all tables in the PostgreSQL database. Every SELECT, "
    "INSERT, UPDATE, and DELETE operation is subject to RLS policies that enforce user ownership "
    "and role constraints at the database level — independent of the application layer. This "
    "provides defense-in-depth: even if application-layer checks were bypassed, the database "
    "would prevent unauthorized data access."
)
bullet("Students can only read their own decisions, skill progress, and profile data.")
bullet("Teachers can read data of students in their own classes only.")
bullet("Administrators can read all data for their role scope.")
bullet("Published scenarios are readable by all authenticated users; draft scenarios only by admins.")

heading("11.3 OWASP Compliance", level=2)
simple_table(
    ["OWASP Category", "FeedWise Mitigation"],
    [
        ["A01: Broken Access Control",     "RBAC + database RLS on all tables"],
        ["A02: Cryptographic Failures",    "All data in transit via HTTPS/TLS; Supabase encrypts at rest"],
        ["A03: Injection",                 "Parameterized queries via Supabase client; Pydantic input validation"],
        ["A04: Insecure Design",           "Role-based architecture reviewed at design phase"],
        ["A05: Security Misconfiguration", "Environment-separated secrets; no default credentials"],
        ["A06: Vulnerable Components",     "Dependency pinning + automated Dependabot alerts"],
        ["A07: Auth Failures",             "JWT verification on every request; token expiry enforced"],
        ["A09: Logging Failures",          "Structured audit_log table; request ID logging throughout"],
        ["A10: SSRF",                      "External URL validation in evidence submission; allowlist enforced"],
    ],
    col_widths=[2.8, 3.8]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — BUSINESS PERSPECTIVE
# ═══════════════════════════════════════════════════════════════════════════════
heading("12. Business Perspective", level=1)
add_hr()

heading("12.1 Revenue Model", level=2)
body(
    "FeedWise is built on a 'freemium for learners, institutional for organizations' model. "
    "Core media literacy learning remains permanently free for youth. Revenue is generated "
    "through institutional partnerships and premium features that enhance the educational "
    "ecosystem without gating individual learning."
)

simple_table(
    ["Revenue Stream", "Target Customers", "Estimated Value"],
    [
        ["School/University Licenses",    "K-12 schools, colleges, universities", "USD 500–2,000/yr per institution"],
        ["NGO & Government Partnerships", "UNESCO, UNICEF, national governments", "Grant-based, USD 50K–500K projects"],
        ["Corporate CSR Programs",        "Tech companies, media organizations", "USD 10K–100K partnerships"],
        ["Teacher Professional Development", "Education ministries, teacher training institutes", "USD 100–500/teacher/yr"],
        ["API Access (Enterprise)",       "EdTech platforms, research institutions", "USD 2,000–20,000/yr"],
        ["Research Data Partnerships",    "Universities, MIL research centers", "Revenue-sharing agreements"],
        ["White-Label Licensing",         "National governments, large NGOs", "USD 50K–200K per deployment"],
    ],
    col_widths=[2.2, 2.0, 2.4]
)

heading("12.2 Market Opportunity", level=2)
body(
    "The global EdTech market was valued at USD 254 billion in 2023 and is projected to reach "
    "USD 605 billion by 2027. The MIL and digital literacy sub-segment is growing at 22% CAGR "
    "driven by government mandates, UNESCO initiatives, and corporate digital responsibility programs."
)
body(
    "FeedWise's primary market is the 1.8 billion young people aged 15–24 globally, with "
    "immediate focus on South Asia (500 million+ youth), where digital adoption is rapid "
    "but media literacy infrastructure is minimal. The platform is uniquely positioned to "
    "serve this market with a multilingual, low-bandwidth-friendly, mobile-first design."
)

heading("12.3 Sustainability Plan", level=2)
simple_table(
    ["Phase", "Timeline", "Key Activity", "Funding Source"],
    [
        ["Phase 1", "Year 1", "Free MVP launch; youth user acquisition", "UNESCO prize, founding grants"],
        ["Phase 2", "Year 1–2", "School partnership pilots; teacher tools", "Institutional licenses"],
        ["Phase 3", "Year 2", "NGO & government partnerships; localization", "Development grants"],
        ["Phase 4", "Year 2–3", "International expansion; API ecosystem", "Enterprise licenses"],
        ["Phase 5", "Year 3+", "Open MIL scenario marketplace", "Platform revenue share"],
    ],
    col_widths=[1.0, 1.0, 2.6, 2.0]
)

body("Core principle: Youth media literacy learning remains free forever.", bold=True, color=DEEP_BLUE)

heading("12.4 Partnership Strategy", level=2)
bullet("UNESCO / UNAOC: Formal partnership for MIL content validation and global distribution.")
bullet("Media Organizations: BBC Media Action, Reuters Foundation, Agence France-Presse for scenario content and professional review.")
bullet("Academic Institutions: Tribhuvan University, IIT, and regional universities for research partnerships and student engagement.")
bullet("Government: National education ministries for curriculum integration and national deployment licenses.")
bullet("Tech Industry: Google.org, Meta for Social Good, Microsoft Philanthropies for CSR funding and infrastructure support.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — CREATIVITY & INNOVATION
# ═══════════════════════════════════════════════════════════════════════════════
heading("13. Creativity & Innovation", level=1)
add_hr()

body(
    "FeedWise is not an incremental improvement on existing media literacy tools. It represents "
    "a fundamentally new paradigm — one that reimagines how MIL education is delivered, "
    "experienced, and measured. This section documents the creative design decisions and "
    "innovations that set FeedWise apart."
)

heading("13.1 Creative Design Principles", level=2)
body(
    "Every creative decision in FeedWise flows from a single insight: young people learn by "
    "doing, not by being told. This led to five foundational creative principles:"
)
simple_table(
    ["Creative Principle", "Design Expression in FeedWise"],
    [
        ["Simulation Over Theory",
         "A realistic, consequence-modelled social media environment replaces textbook lessons — "
         "every action teaches rather than every page"],
        ["Consequences as Curriculum",
         "The Consequence Engine makes the cost of poor information hygiene tangible and immediate, "
         "not abstract — learners feel the impact of their decisions"],
        ["Evidence Over Verdicts",
         "TrustLens never says 'this is fake.' It shows evidence and asks 'what do YOU conclude?' — "
         "teaching the process of reasoning, not dependence on verdicts"],
        ["AI Assists, Humans Decide",
         "AI analysis is one input among many — never the final authority — the creative inverse "
         "of most AI fact-checking tools that replace human judgment"],
        ["Progress You Can See",
         "The 5-dimension radar chart transforms abstract skill growth into a visible, living "
         "portrait of a learner's MIL competency across five distinct dimensions"],
    ],
    col_widths=[2.2, 4.4]
)

heading("13.2 Key Innovations", level=2)
bullet("TrustLens Investigation Panel: The first structured, multi-dimension evidence investigation tool designed specifically for youth — combining source credibility, author identity, date verification, evidence quality, and AI language signal analysis into a single guided workflow.")
bullet("Consequence Engine: A real-time impact simulator that translates information decisions into measurable simulated outcomes — reach numbers, credibility scores, community effects. No equivalent exists in any existing MIL platform.")
bullet("AI-Assisted, Human-Reviewed Pipeline: AI drafts explanations and flags linguistic signals, but every scenario is reviewed and approved by a qualified MIL professional before publication — combining the scale of AI with the integrity of expert oversight.")
bullet("5-Dimension Skill Radar: A continuous, multi-axis competency visualization mapping growth across Source Verification, Evidence Evaluation, AI Literacy, Bias Detection, and Digital Safety — providing a holistic, actionable portrait of MIL development.")
bullet("Newsroom Mode (Editorial Perspective Shift): An innovative role-reversal exercise — students become editors, experiencing the pressures and ethical trade-offs of responsible journalism and building empathy alongside critical thinking skills.")
bullet("Community Scenario Ecosystem: A self-sustaining platform where advanced learners contribute real-world claims for peer investigation, transforming passive learners into active MIL community contributors and content creators.")

heading("13.3 FeedWise vs. Existing Landscape", level=2)
simple_table(
    ["What Exists Today",                              "FeedWise Innovation"],
    [
        ["Fact-checkers issue binary TRUE / FALSE",         "FeedWise teaches the evaluation process — never issues a verdict"],
        ["Educational apps deliver content passively",      "FeedWise puts learners inside a simulated information environment"],
        ["Games focus on entertainment value",              "Every FeedWise mechanic maps to a measurable MIL learning outcome"],
        ["AI tools replace human judgment",                 "FeedWise AI augments investigation; MIL experts review every result"],
        ["Individual-focused learning tools",               "FeedWise integrates classroom, community, and professional layers"],
        ["Progress as scores or points",                    "FeedWise tracks 5-dimensional skill development, not just correct answers"],
        ["English-only, Western-context content",           "FeedWise is multilingual-by-design with culturally adapted libraries"],
        ["One-time awareness campaigns",                    "FeedWise builds a self-sustaining, community-powered scenario ecosystem"],
    ],
    col_widths=[3.0, 3.6]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 14 — FEASIBILITY ASSESSMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading("14. Feasibility Assessment", level=1)
add_hr()

body(
    "FeedWise is not a concept paper — it is a working platform with a clear technical foundation, "
    "a realistic cost structure, and a validated market need. This section provides an honest "
    "assessment of technical, financial, operational, and market feasibility."
)

heading("14.1 Technical Feasibility", level=2)
body(
    "Every technology used in FeedWise is production-proven and widely deployed at scale. "
    "No experimental or unproven technology is required for the platform to function."
)
simple_table(
    ["Component",              "Technology",                   "Feasibility Evidence"],
    [
        ["Mobile + Web App",       "Flutter (Dart) 3.27+",          "Powers BMW, eBay, Google Pay apps globally"],
        ["Admin Portal",           "Flutter Web",                   "Same codebase as user app; proven cross-platform"],
        ["Backend API",            "FastAPI (Python) 0.115+",       "Used by Microsoft, Uber, Netflix for high-throughput APIs"],
        ["Database",               "PostgreSQL via Supabase",        "Managed at scale; enterprise SLA available; open-source"],
        ["Authentication",         "Supabase Auth (JWT/OAuth2)",     "Industry standard; battle-tested security model"],
        ["AI/NLP",                 "OpenAI / Gemini API",            "Stable, documented; fallback model switching supported"],
        ["Hosting",                "Render.com + Firebase Hosting",  "Free-tier viable; auto-deploys from GitHub on push"],
    ],
    col_widths=[1.8, 2.0, 2.8]
)
body("Current Development Status (as of August 2026):", bold=True, color=DEEP_BLUE)
bullet("Backend API: Fully scaffolded with FastAPI; all core routes implemented; 10+ automated tests passing.")
bullet("Database: Complete 25-table PostgreSQL schema with Row Level Security policies and seed data deployed.")
bullet("Flutter User App: Core navigation, feed, TrustLens investigation, decision, and profile screens implemented.")
bullet("Flutter Admin Studio: Scenario CMS and teacher dashboard architecture scaffolded and in development.")
bullet("AI Engine: Text signal analysis and explanation generation module integrated and tested with mock scenarios.")

heading("14.2 Financial & Operational Feasibility", level=2)
simple_table(
    ["Cost Category",             "Year 1 Estimate",       "Mitigation Strategy"],
    [
        ["Infrastructure (Hosting)",   "USD 0–200 / month",   "Render.com + Supabase free tiers; costs scale with users"],
        ["Domain & SSL Certificate",   "USD 15 / year",       "Standard annual cost"],
        ["AI API Calls",               "USD 50–500 / month",  "Caching reduces repeat analysis costs; per-call billing"],
        ["Development Team",           "Voluntary (students)","All 5 members contributing on a voluntary basis for submission"],
        ["MIL Content Review",         "Volunteer reviewers", "UNESCO partnership and educator volunteer program"],
        ["Total Year 1 Estimate",      "USD 600–8,400",       "Coverable by UNESCO prize, seed grant, or university support"],
    ],
    col_widths=[2.4, 1.6, 2.6]
)
body(
    "The freemium model ensures the platform operates at near-zero cost during user acquisition "
    "while building the base needed to justify institutional licensing revenue from Year 2 onwards."
)

heading("14.3 Market Feasibility", level=2)
simple_table(
    ["Feasibility Dimension",  "Evidence"],
    [
        ["Market Demand",          "UNESCO 2026 MIL Challenge validates global institutional interest; 1.8B youth in target demographic"],
        ["Problem Urgency",        "AI-generated misinformation growing exponentially; government mandates for digital literacy increasing globally"],
        ["Competitive Gap",        "No existing platform combines simulation + consequence engine + teacher integration + AI assistance in one tool"],
        ["Partnership Readiness",  "UNESCO, BBC Media Action, Reuters Foundation, and Google.org all have active MIL funding programs"],
        ["Revenue Validation",     "School licensing model validated by EdTech tools: Newsela, CommonSense Media, Quill.org"],
        ["Scale Path",             "Software scales without proportional cost increase; global deployment identical to local deployment technically"],
    ],
    col_widths=[2.4, 4.2]
)

heading("14.4 Risk Assessment & Mitigation", level=2)
simple_table(
    ["Risk",                               "Likelihood", "Impact", "Mitigation Strategy"],
    [
        ["AI API costs exceed budget",          "Medium", "Medium", "Aggressive caching; local fallback model; usage caps per user/day"],
        ["Low initial user adoption",           "Medium", "High",   "School partnerships provide captive initial user base from launch"],
        ["Content quality issues",              "Low",    "High",   "Human MIL expert mandatory review before any scenario is published"],
        ["Data privacy breach",                 "Low",    "High",   "Supabase RLS + JWT auth + minimal PII collection by design"],
        ["Platform moderation at scale",        "Medium", "Medium", "Tiered moderation: AI flagging first, human review pipeline second"],
        ["Team bandwidth (student developers)", "High",   "Medium", "Phased roadmap; MVP scope carefully limited to core learning loop"],
    ],
    col_widths=[2.6, 0.9, 0.7, 2.4]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 15 — FUTURE DEVELOPMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading("15. Future Development", level=1)
add_hr()

heading("15.1 Phase Roadmap", level=2)
simple_table(
    ["Phase", "Name", "Key Features"],
    [
        ["Phase 1 (MVP)", "Foundation", "Core feed, TrustLens, Decision Engine, basic Academy, skill radar"],
        ["Phase 2", "Classroom Integration", "Teacher dashboard, class management, assignments, analytics"],
        ["Phase 3", "Game Simulation", "Newsroom Zero, Creator Mode, scenario leaderboards"],
        ["Phase 4", "Community", "Community submissions, peer review, user-generated challenges"],
        ["Phase 5", "Localization", "Nepali, Hindi, Bengali, Arabic — full translated scenario libraries"],
        ["Phase 6", "AI Enhancement", "Adaptive difficulty, personalized learning paths, AI tutor"],
        ["Phase 7", "API Ecosystem", "Open API for third-party EdTech integrations; LTI support"],
        ["Phase 8", "Research Platform", "Academic research portal, anonymized data access, impact measurement"],
    ],
    col_widths=[1.2, 1.8, 3.6]
)

heading("15.2 Planned Features", level=2)
bullet("Adaptive Difficulty Engine: Dynamically adjusts scenario difficulty based on real-time skill progression analysis.")
bullet("AI Tutor (FeedBot): A conversational AI tutor that guides students through investigations and explains MIL concepts in plain language.")
bullet("Deepfake Detection Lab: A specialized module using visual and audio analysis tools to teach identification of synthetic media.")
bullet("Multiplayer Newsroom: A collaborative mode where teams of 3–5 students work together on editorial decisions in real time.")
bullet("Open Scenario API: Allow verified educators globally to submit scenarios through a standardized API, growing the content library at scale.")
bullet("LTI Integration: Learning Tools Interoperability (LTI) compliance for seamless integration with LMS platforms (Moodle, Canvas, Google Classroom).")
bullet("Offline Mode: Full offline scenario cache for users in low-connectivity environments, syncing progress when connected.")
bullet("Parental Dashboard: Guardian-accessible summary view of student MIL development for younger learners (under 16).")

heading("15.3 Research & Impact Measurement", level=2)
body(
    "FeedWise is designed to generate research-quality learning outcome data. Future development "
    "includes a formal Research Portal with anonymized dataset access for academic institutions "
    "studying digital literacy, misinformation resistance, and AI literacy development in youth. "
    "The platform will partner with universities to conduct longitudinal studies on MIL skill "
    "development, with the goal of establishing FeedWise's methodology as an evidence-based "
    "standard for MIL education globally."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 16 — CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
heading("16. Conclusion", level=1)
add_hr()

body(
    "FeedWise represents a new paradigm in Media and Information Literacy education. At a "
    "moment when the information environment has never been more complex, more dangerous, or "
    "more consequential, FeedWise offers a practical, proven-methodology response: teach young "
    "people not what to think, but how to think."
)
body(
    "By combining a realistic social media simulation with a structured evidence investigation "
    "framework, a consequence engine that makes the impact of information decisions tangible, "
    "an AI engine that assists without replacing critical thinking, and a classroom integration "
    "system that connects individual student learning to professional educator oversight, "
    "FeedWise creates an educational experience that is uniquely effective."
)
body(
    "The platform is technically grounded — built on proven, scalable technologies with a "
    "clear deployment path, a sustainable business model, and an open-source ethos that "
    "invites collaboration. It is educationally sound — every feature is designed around "
    "measurable MIL learning outcomes. And it is deeply human — in a world where AI generates "
    "more content every day, FeedWise insists that human judgment, human curiosity, and "
    "human critical thinking remain at the center of the information ecosystem."
)
body(
    "The team behind FeedWise — five student developers from Nepal — built this platform "
    "because they live inside this problem. They navigate the same feeds, encounter the same "
    "misinformation, and face the same AI-era challenges as the students FeedWise is designed "
    "to serve. That lived experience is embedded in every design decision, every scenario, "
    "and every line of code."
)
body(
    "FeedWise is a UNESCO MIL Challenge submission. But it is also a long-term commitment. "
    "The platform will continue to grow — more scenarios, more languages, more users, more "
    "partnerships — because the problem it addresses will not go away. If anything, it will "
    "grow more urgent with every advance in generative AI."
)

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_final = p_final.add_run('"In the age of AI, the most important skill is knowing how to think — not what to think."')
set_font(r_final, "Calibri", 14, italic=True, color=DEEP_BLUE)
p_final.paragraph_format.space_before = Pt(16)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 17 — REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
heading("17. References", level=1)
add_hr()

references = [
    "UNESCO (2023). Global Framework for Media and Information Literacy. Paris: UNESCO Publishing.",
    "UNESCO (2022). MIL Curriculum for Teachers. Paris: UNESCO.",
    "Wardle, C., & Derakhshan, H. (2017). Information Disorder: Toward an interdisciplinary framework for research and policy making. Council of Europe.",
    "EU DisinfoLab (2023). State of Disinformation in the AI Era. Brussels: EU DisinfoLab.",
    "Lewandowsky, S., Ecker, U. K. H., & Cook, J. (2017). Beyond Misinformation: Understanding and Coping with the Post-Truth Era. Journal of Applied Research in Memory and Cognition, 6(4), 353-369.",
    "Shu, K., Sliva, A., Wang, S., Tang, J., & Liu, H. (2017). Fake News Detection on Social Media: A Data Mining Perspective. ACM SIGKDD Explorations Newsletter, 19(1), 22-36.",
    "Flutter Development Team (2024). Flutter Documentation. Google LLC. https://flutter.dev/docs",
    "FastAPI Documentation (2024). Sebastián Ramírez. https://fastapi.tiangolo.com",
    "Supabase Documentation (2024). Supabase Inc. https://supabase.com/docs",
    "OWASP Foundation (2021). OWASP Top Ten 2021. https://owasp.org/www-project-top-ten/",
    "UNESCO Global MIL Week (2026). UNESCO MIL Youth Challenge 2026 Guidelines. Paris: UNESCO.",
    "Google DeepMind (2024). Gemini 2.0 Technical Report. Google LLC.",
    "OpenAI (2024). GPT-4o Technical Overview. OpenAI.",
    "Global EdTech Market Report (2023). HolonIQ Education Intelligence. https://www.holoniq.com",
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"[{i}]  {ref}")
    set_font(r, "Calibri", 10.5, color=DARK_GREY)
    p.paragraph_format.space_after = Pt(5)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX A — SCREENSHOTS PLACEHOLDER
# ═══════════════════════════════════════════════════════════════════════════════
heading("Appendix A: Interface Screenshots", level=1)
add_hr()

body(
    "The following screenshots were captured from the FeedWise application during development "
    "and testing. Screenshots will be inserted manually by the team before final submission."
)

screenshots = [
    ("A.1", "Splash Screen & Onboarding", "FeedWise launch screen with 3-slide onboarding flow"),
    ("A.2", "Social Feed (Home Screen)", "Main feed showing scenario cards with engagement metrics and trending badges"),
    ("A.3", "TrustLens Investigation Panel", "Full investigation panel showing source, evidence, and language signal analysis"),
    ("A.4", "Decision Screen", "Four-option decision interface (Share / Verify / Report / Ignore)"),
    ("A.5", "Consequence Card", "Post-decision consequence visualization with reach numbers and credibility delta"),
    ("A.6", "MIL Lesson Screen", "Lesson delivery screen with key takeaway and follow-up quiz"),
    ("A.7", "Skill Radar Chart", "5-dimension MIL competency radar visualization"),
    ("A.8", "Badge Collection", "User achievement badges and unlocking progress"),
    ("A.9", "MIL Academy", "Academy module browser with skill dimension categories"),
    ("A.10", "Newsroom Mode", "Editorial simulation interface for Newsroom Zero scenarios"),
    ("A.11", "Teacher Dashboard", "Class overview with student skill analytics and assignment tools"),
    ("A.12", "Admin Scenario CMS", "Scenario creation and review pipeline in FeedWise Studio"),
]

simple_table(
    ["Ref", "Screen Name", "Description"],
    screenshots,
    col_widths=[0.5, 2.2, 4.0]
)

body(
    "\n[Screenshot images to be inserted by team before submission. "
    "Placeholder spaces reserved for 12 interface screenshots.]",
    italic=True, color=MID_GREY
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/home/upendra-singh-dhami/Documents/feedwise/documentation/unescoproject.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
